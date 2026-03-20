import subprocess, sys
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

OUTPUT = "/Users/jairfahl/Downloads/c-level/C-Level_Mentor_Roadmap.docx"

# ── Helpers ─────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text="", size=11, bold=False, italic=False,
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading(doc, text, level=1):
    colors = {1: (30, 58, 95), 2: (30, 58, 95), 3: (60, 90, 130)}
    sizes  = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=sizes[level], bold=True, color=colors[level])
    if level == 1:
        # linha horizontal sob o título
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1e3a5f")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def add_bullet(doc, text, indent_level=0, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + indent_level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # azul escuro no header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1e3a5f")
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        fill = "f0f4f8" if ri % 2 == 0 else "ffffff"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

def add_clevel_section(doc, role, subtitle, domain, decisions, frameworks, game_theory, differentials, timeline):
    add_heading(doc, f"{role} — {subtitle}", level=2)
    add_paragraph(doc, f"Período estimado: {timeline}", size=10, italic=True, color=(100, 100, 100))
    items = [
        ("Domínio", domain),
        ("Decisões-tipo", decisions),
        ("Frameworks", frameworks),
        ("Teoria dos Jogos", game_theory),
        ("Diferenciais vs CFO Mentor", differentials),
    ]
    for label, content in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(f"{label}: ")
        set_font(r1, size=11, bold=True)
        r2 = p.add_run(content)
        set_font(r2, size=11)
    doc.add_paragraph()

# ── Documento ────────────────────────────────────────────────────────────────

doc = Document()

# Margens
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Fonte padrão
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── CAPA ─────────────────────────────────────────────────────────────────────
add_paragraph(doc, "", space_before=60)
add_paragraph(doc, "C-LEVEL MENTOR", size=28, bold=True,
              color=(30, 58, 95), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph(doc, "Roadmap de Produto", size=18, bold=False,
              color=(30, 58, 95), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_paragraph(doc, "Visão de Futuro e Expansão da Plataforma para o C-Suite", size=13,
              italic=True, color=(90, 110, 140), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)
add_paragraph(doc, "Versão 1.0  ·  Março 2026  ·  Confidencial", size=10,
              color=(130, 130, 130), align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ── POSICIONAMENTO DO PRODUTO ─────────────────────────────────────────────────
add_heading(doc, "Posicionamento do Produto", level=1)
add_paragraph(doc,
    "O C-Level Mentor não é um chatbot nem um assistente conversacional. É uma plataforma de "
    "agentes especializados — sistemas capazes de executar fluxos de trabalho executivos completos "
    "de forma estruturada, combinando raciocínio de LLM com regras de negócio determinísticas.",
    space_after=6)

add_heading(doc, "Os três componentes de cada agente C-Level", level=2)
componentes = [
    ("Modelo", "Claude claude-opus-4-5 — o cérebro que raciocina sobre premissas, riscos, frameworks e contexto financeiro/operacional de cada decisão."),
    ("Ferramentas", "FastAPI endpoints + PostgreSQL + Redis — APIs que permitem ao agente persistir decisões, consultar histórico, acionar frameworks e registrar audit trail completo."),
    ("Instruções", "Engine Determinística (PromptBuilder + impact_scorer + framework_selector + game_theory_activator) — diretrizes explícitas que definem como o agente raciocina, quais frameworks aplica e quais limites respeita."),
]
add_table(doc,
    headers=["Componente", "Implementação no C-Level Mentor"],
    rows=componentes,
    col_widths=[3.5, 12.5])
doc.add_paragraph()

add_heading(doc, "Guardrails — Defesa em Camadas", level=2)
guardrails = [
    "State machine linear — impede transições inválidas (HTTP 409); o agente nunca pula etapas.",
    "Fallback determinístico — qualquer falha do LLM aciona análise estruturada; o agente nunca para.",
    "Audit trail completo — cada ação do agente é registrada com timestamp, ator e payload.",
    "Intervenção humana obrigatória — o agente recomenda; o executivo decide (RECOMMENDED → DECIDED).",
    "Revisão pós-decisão — ciclo de aprendizado estruturado no estado UNDER_REVIEW → CLOSED.",
]
for g in guardrails:
    add_bullet(doc, g)
doc.add_paragraph()

add_heading(doc, "CEO Orchestrator — O Agente-Mestre", level=2)
add_paragraph(doc,
    "O CEO Orchestrator, previsto para Q2–Q3 2027, é o agente de fluxo completo descrito na "
    "literatura de IA agentica: consolida os outputs de todos os módulos C-Level (ferramentas), "
    "aplica raciocínio estratégico sobre conflitos entre recomendações (modelo) e opera dentro "
    "das diretrizes de governança corporativa do board (instruções). É a camada que transforma "
    "a plataforma de agentes especializados em um Board Decision Support System autônomo.",
    space_after=12)

# ── 1. SUMÁRIO EXECUTIVO ─────────────────────────────────────────────────────
add_heading(doc, "1. Sumário Executivo", level=1)
add_paragraph(doc,
    "O C-Level Mentor é uma plataforma de apoio à decisão executiva que combina três pilares "
    "inéditos: IA Generativa (Claude Anthropic), Engine Determinística (frameworks financeiros e "
    "operacionais codificados) e Teoria dos Jogos (ativada automaticamente quando há agentes "
    "externos na decisão). O módulo-âncora — o CFO Mentor — encontra-se plenamente operacional, "
    "com stack FastAPI + PostgreSQL + React 18 e cobertura de testes de 99%.",
    space_after=6)
add_paragraph(doc,
    "Este roadmap descreve a evolução natural da plataforma para cobrir todos os papéis do "
    "C-Suite — COO, CIO, CMO, CRO, CHRO, CLO/CGO e CISO — culminando no CEO Orchestrator, "
    "camada de integração que transforma o produto em um Board Decision Support System completo. "
    "Cada novo módulo reutiliza a mesma arquitetura de state machine, audit trail e LLM layer; "
    "apenas os domínios, frameworks e métricas de risco variam por papel executivo.",
    space_after=10)

# ── 2. STATUS ATUAL — CFO MENTOR ─────────────────────────────────────────────
add_heading(doc, "2. Status Atual — CFO Mentor", level=1)
add_paragraph(doc, "Fases concluídas:", bold=True)
fases = [
    ("P-01", "Core / Main", "Concluído"),
    ("P-02", "ORM Models", "Concluído"),
    ("P-03", "Schemas Pydantic v2", "Concluído — 73 testes"),
    ("P-04", "Engine Determinística (state machine, impact scorer, framework selector, game theory, audit)", "Concluído — 108 testes"),
    ("P-05", "Camada LLM (prompt builder, client, parser, cache, fallback)", "Concluído — 76 testes"),
    ("P-06", "API Routers (cases, state machine, audit)", "Concluído — 43 testes de integração"),
    ("P-07", "Integração end-to-end + smoke tests", "Concluído — 320 testes, 99% cobertura"),
    ("P-FE", "Frontend React 18 + Vite + Tailwind CSS", "Concluído — http://localhost:3000"),
]
add_table(doc,
    headers=["Fase", "Descrição", "Status"],
    rows=fases,
    col_widths=[1.5, 10.0, 4.5])
doc.add_paragraph()

add_paragraph(doc, "Cobertura das Top-5 Técnicas de Decisão Executiva:", bold=True)
tecnicas = [
    ("Planejamento de Cenários", "✅ Nativo — framework scenario_analysis ativado para impacto ≥ 4; LLM gera scenario_summary"),
    ("IA + BRMS", "✅ Nativo — Engine determinística (BRMS) + Claude claude-opus-4-5 (IA Generativa)"),
    ("Tomada de Decisão Baseada em Dados", "🔶 Parcial — exposição financeira, premissas e riscos estruturados; KPIs em tempo real: backlog"),
    ("Metodologias Ágeis", "🔶 Parcial — botão 'Reanalisar' implementado; loops iterativos completos: backlog"),
    ("Decisão Colaborativa", "🔲 Backlog — divergence_flag CFO vs LLM implementado; multi-stakeholder no UNDER_REVIEW: backlog"),
]
add_table(doc,
    headers=["Técnica", "Cobertura Atual"],
    rows=tecnicas,
    col_widths=[5.0, 11.0])
doc.add_paragraph()

# ── 3. PRÓXIMOS CAPÍTULOS — EXTENSÕES CFO ────────────────────────────────────
add_heading(doc, "3. Próximos Capítulos — Extensões do CFO Mentor", level=1)
ext_rows = [
    ("3.0", "Multi-Framework Selection — executivo vê e customiza os métodos antes da análise (mín. 1, máx. 4); LLM combina e sintetiza recomendação integrada", "Todas", "Alta ★", "Q2 2026"),
    ("3.1", "KPIs financeiros em tempo real integrados ao case como inputs automáticos de contexto", "Data-Driven", "Alta", "Q2 2026"),
    ("3.2", "Delphi Digital no UNDER_REVIEW — aprovação multi-stakeholder com rodadas estruturadas de consenso e registro de convergência", "Colaborativa + Delphi", "Alta", "Q2 2026"),
    ("3.3", "Agile Decision Loop — reanálise iterativa com botão 'Reanalisar' em RECOMMENDED", "Ágil", "Feito ✅", "Q1 2026"),
    ("3.4", "Dashboard analítico de decisões históricas — acurácia de forecast, taxa de realização de riscos, heurísticas geradas", "Data-Driven", "Média", "Q3 2026"),
    ("3.5", "Integração com ERPs (SAP/Oracle) para ingestão automática de dados financeiros", "Data-Driven", "Futura", "2027"),
]
add_table(doc,
    headers=["#", "Feature", "Técnica Apoiada", "Prioridade", "Período"],
    rows=ext_rows,
    col_widths=[0.8, 8.5, 3.2, 2.0, 1.8])
doc.add_paragraph()

add_heading(doc, "Biblioteca de Métodos — 11 Frameworks Disponíveis", level=2)
add_paragraph(doc,
    "A partir da feature 3.0 (Multi-Framework Selection), o executivo poderá combinar "
    "até 4 métodos por análise. O LLM aplica cada framework como uma lente complementar "
    "sobre o mesmo case e sintetiza uma recomendação integrada — prática padrão em "
    "consultoria estratégica (McKinsey, BCG). A engine sugere a combinação ideal; "
    "o executivo confirma ou ajusta.",
    space_after=6)

metodos_rows = [
    ("pdca", "PDCA", "Melhoria contínua, ajustes de processo e budget", "CFO, COO"),
    ("scenario_analysis", "Análise de Cenários", "Decisões com incerteza e múltiplos futuros possíveis", "CFO, CIO, CEO"),
    ("game_theory", "Teoria dos Jogos", "Decisões com agentes externos (negociação, concorrência)", "CFO, CRO, CEO"),
    ("trade_off", "Trade-Off", "Decisões binárias com ganhos e perdas explícitos", "CFO, COO"),
    ("risk_matrix", "Matriz de Riscos", "Priorização de riscos por probabilidade × impacto", "CFO, CISO, CLO"),
    ("capital_allocation", "Alocação de Capital", "Distribuição de recursos entre múltiplos projetos", "CFO, CEO"),
    ("decision_matrix", "Matriz de Decisão ★ novo", "Avaliação multi-critério ponderada entre alternativas", "CFO, CIO, COO"),
    ("cost_benefit_analysis", "Análise Custo-Benefício ★ novo", "Quantificação de NPV, IRR e payback de iniciativas", "CFO, CIO, CRO"),
    ("decision_tree", "Árvore de Decisão ★ novo", "Mapeamento probabilístico de cenários com EMV", "CFO, CEO, CRO"),
    ("swot_analysis", "Análise SWOT/FOFA ★ novo", "Contexto competitivo e estratégico da decisão", "CMO, CEO, CRO"),
    ("delphi_method", "Método Delphi ★ novo", "Consenso estruturado multi-stakeholder (UNDER_REVIEW)", "Todos os C-Levels"),
]
add_table(doc,
    headers=["Código", "Framework", "Quando Usar", "Módulos C-Level"],
    rows=metodos_rows,
    col_widths=[3.2, 3.8, 6.5, 2.5])
doc.add_paragraph()

# ── 4. EXPANSÃO C-LEVEL ───────────────────────────────────────────────────────
add_heading(doc, "4. Expansão C-Level — Módulos Futuros", level=1)
add_paragraph(doc,
    "Cada módulo compartilha a mesma arquitetura de state machine, audit trail, camada LLM e "
    "frontend React — apenas os domínios, frameworks, métricas e tipos de decisão variam. "
    "A expansão é, portanto, predominantemente de configuração e conteúdo, não de engenharia.",
    space_after=8)

c_levels = [
    {
        "role": "4.1  COO — Chief Operating Officer",
        "subtitle": "Excelência Operacional",
        "domain": "Operações, supply chain, capacidade produtiva, SLAs, processos.",
        "decisions": "Otimização de processos, negociação com fornecedores, planejamento de capacidade, decisões de outsourcing, gestão de gargalos.",
        "frameworks": "Lean, Six Sigma, TOC (Theory of Constraints), SIPOC, OKR operacional, RACI.",
        "game_theory": "Negociação com fornecedores em oligopólio (Stackelberg), competição de capacidade entre unidades de negócio.",
        "differentials": "KPIs operacionais como inputs (OEE, lead time, fill rate, nível de serviço); simulador de gargalos produtivos.",
        "timeline": "Q2 2026",
    },
    {
        "role": "4.2  CIO — Chief Information Officer",
        "subtitle": "Estratégia de Tecnologia",
        "domain": "Estratégia de TI, transformação digital, cibersegurança, arquitetura de sistemas, FinOps.",
        "decisions": "Build vs buy, seleção de fornecedor tecnológico, priorização de portfólio IT, gestão de licenças, migração para cloud.",
        "frameworks": "TOGAF, COBIT, ITIL, DevOps/DORA, FinOps, Zero Trust.",
        "game_theory": "Seleção de plataformas com lock-in e custos de switching (Bertrand com diferenciação); negociação de contratos de licença de longo prazo.",
        "differentials": "TCO/ROI de iniciativas tecnológicas; risk matrix de cibersegurança como input de risco; análise de dependência tecnológica.",
        "timeline": "Q3 2026",
    },
    {
        "role": "4.3  CMO — Chief Marketing Officer",
        "subtitle": "Crescimento e Marca",
        "domain": "Marca, geração de demanda, experiência do cliente, go-to-market, pricing.",
        "decisions": "Alocação de budget de mídia, lançamento de produto, expansão geográfica, mudança de posicionamento, decisão de pricing.",
        "frameworks": "Growth loops, Attribution (MMM/MTA), CAC/LTV, Blue Ocean Strategy, JTBD (Jobs-to-be-Done).",
        "game_theory": "Guerras de preço (Nash equilibrium de pricing), timing de lançamento vs concorrente (first-mover advantage), disputa por share of voice em mídia programática.",
        "differentials": "Funil de conversão e elasticidade de preço como inputs; NPS/CSAT como indicadores de risco de churn.",
        "timeline": "Q4 2026",
    },
    {
        "role": "4.4  CRO — Chief Revenue Officer",
        "subtitle": "Receita e Expansão Comercial",
        "domain": "Receita, precificação, estratégia comercial, expansão de mercado, incentivos de vendas.",
        "decisions": "Mudança de pricing, design de comissionamento, expansão para novo segmento, aprovação de contrato enterprise, política de descontos.",
        "frameworks": "RevOps, Value-Based Pricing, MEDDIC, análise de NRR/ARR/MRR, cohort analysis.",
        "game_theory": "Negociação de contratos de grande porte (leilão de Vickrey); competição de pricing em mercados com poucos players.",
        "differentials": "Métricas SaaS como inputs nativos (Churn, MRR, ARR, LTV); pipeline forecast como contexto de risco.",
        "timeline": "Q4 2026",
    },
    {
        "role": "4.5  CHRO — Chief Human Resources Officer",
        "subtitle": "Pessoas, Cultura e Talentos",
        "domain": "Pessoas, talentos, cultura organizacional, remuneração, reestruturação.",
        "decisions": "Revisão salarial e bandas, succession planning, decisão de layoff, política de benefícios, programa DEI, restruturação organizacional.",
        "frameworks": "9-Box Grid, Total Rewards, OKR de pessoas, Kotter Change Model, modelo de maturidade DEI.",
        "game_theory": "Retenção de talentos escassos vs custo (barganha bilateral); negociação salarial com candidatos em mercado aquecido; competição por atração de executivos.",
        "differentials": "eNPS, turnover voluntário e absenteísmo como inputs de risco; assessment de sucessão como contexto decisório.",
        "timeline": "Q1 2027",
    },
    {
        "role": "4.6  CLO/CGO — Chief Legal / Governance Officer",
        "subtitle": "Jurídico, Compliance e ESG",
        "domain": "Jurídico, compliance regulatório, ESG, governança corporativa.",
        "decisions": "Estratégia de litígio, investimento em compliance regulatório, definição de commitments ESG, adequação LGPD/GDPR, gestão de terceiros (due diligence).",
        "frameworks": "GRC (Governance Risk Compliance), ISO 37001 (antissuborno), TCFD (ESG/clima), LGPD/GDPR Maturity Model.",
        "game_theory": "Estratégia sequencial em litígio (Stackelberg — quem cede primeiro); negociação com órgãos reguladores (cooperação vs contestação).",
        "differentials": "Risk registry legal como input; scoring ESG e matriz de materialidade; análise de exposição regulatória setorial.",
        "timeline": "Q1 2027",
    },
    {
        "role": "4.7  CISO — Chief Information Security Officer",
        "subtitle": "Cibersegurança e Privacidade",
        "domain": "Cibersegurança, privacidade de dados, risco tecnológico, gestão de incidentes.",
        "decisions": "Investimento em segurança vs exposição ao risco, resposta a incidentes (pagar ransomware ou não), gestão de risco de terceiros, priorização de patches.",
        "frameworks": "NIST Cybersecurity Framework, ISO 27001, Zero Trust, MITRE ATT&CK, FAIR (Cyber Risk Quantification).",
        "game_theory": "Modelo atacante-defensor (Stackelberg — defensor como líder de Stackelberg); alocação ótima de defesa entre múltiplos vetores de ataque.",
        "differentials": "CVE scoring e CVSS como inputs de risco; quantificação financeira de cyber risk (modelo FAIR); probabilidade de breach por setor.",
        "timeline": "Q2 2027",
    },
    {
        "role": "4.8  CEO — Chief Executive Officer",
        "subtitle": "Estratégia Corporativa — Orchestrator",
        "domain": "Estratégia corporativa, fusões e aquisições, governança, stakeholders, visão de longo prazo.",
        "decisions": "Pivô estratégico, M&A (buy vs build vs partner), estrutura de capital, posicionamento competitivo, decisões de desinvestimento.",
        "frameworks": "Porter's Five Forces, Blue Ocean Strategy, BCG Matrix, Balanced Scorecard, Business Model Canvas, Jobs-to-be-Done.",
        "game_theory": "Negociação de M&A (leilão de Vickrey — valor de reserva); estratégia competitiva de mercado (Nash equilibrium setorial); decisão de entrada/saída de mercado.",
        "differentials": "Papel único de CEO Orchestrator: consolida outputs de todos os módulos C-Level em um painel unificado; detecta conflitos entre recomendações dos C-Levels e propõe mediação via LLM; camada de visão de portfólio estratégico.",
        "timeline": "Q2–Q3 2027",
    },
]

for cl in c_levels:
    add_clevel_section(doc,
        role=cl["role"], subtitle=cl["subtitle"],
        domain=cl["domain"], decisions=cl["decisions"],
        frameworks=cl["frameworks"], game_theory=cl["game_theory"],
        differentials=cl["differentials"], timeline=cl["timeline"])

# ── 5. VISÃO DA PLATAFORMA C-SUITE ───────────────────────────────────────────
add_heading(doc, "5. Visão da Plataforma C-Suite Completa", level=1)
add_paragraph(doc,
    "O CEO Orchestrator é a camada de integração que transforma o conjunto de módulos em um "
    "Board Decision Support System. Cada módulo opera de forma autônoma, mas expõe seus outputs "
    "para o Orchestrator agregar, detectar conflitos e propor síntese estratégica via LLM.",
    space_after=8)

diagrama = (
    "CEO Orchestrator  (Board Decision Support System)\n"
    "│\n"
    "├── CFO Mentor       ·  Financeiro              ← Concluído ✅\n"
    "├── COO Mentor       ·  Operacional             ← Q2 2026\n"
    "├── CIO Mentor       ·  Tecnologia              ← Q3 2026\n"
    "├── CMO Mentor       ·  Marketing               ← Q4 2026\n"
    "├── CRO Mentor       ·  Receita                 ← Q4 2026\n"
    "├── CHRO Mentor      ·  Pessoas                 ← Q1 2027\n"
    "├── CLO/CGO Mentor   ·  Jurídico / Governança   ← Q1 2027\n"
    "├── CISO Mentor      ·  Cibersegurança          ← Q2 2027\n"
    "└── CEO Orchestrator ·  Síntese Estratégica     ← Q2–Q3 2027"
)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(8)
run = p.add_run(diagrama)
run.font.name = "Courier New"
run.font.size = Pt(9)

# ── 6. TIMELINE ───────────────────────────────────────────────────────────────
add_heading(doc, "6. Timeline", level=1)
timeline_rows = [
    ("Q1 2026", "CFO Mentor completo", "Concluído ✅"),
    ("Q2 2026", "Extensões CFO: KPIs em tempo real + multi-stakeholder UNDER_REVIEW · COO Mentor", "Planejado"),
    ("Q3 2026", "CIO Mentor · Dashboard analítico histórico de decisões", "Planejado"),
    ("Q4 2026", "CMO Mentor · CRO Mentor", "Planejado"),
    ("Q1 2027", "CHRO Mentor · CLO/CGO Mentor", "Planejado"),
    ("Q2 2027", "CISO Mentor · CEO Orchestrator (Fase 1 — painel unificado)", "Planejado"),
    ("Q3 2027", "Plataforma C-Suite completa · Board Decision Support System · GA", "Visão"),
]
add_table(doc,
    headers=["Período", "Entregável", "Status"],
    rows=timeline_rows,
    col_widths=[2.0, 11.5, 2.5])
doc.add_paragraph()

# ── 7. PREMISSAS ARQUITETURAIS ────────────────────────────────────────────────
add_heading(doc, "7. Premissas Arquiteturais para Expansão", level=1)
premissas = [
    "State machine, audit trail e camada LLM são reutilizados sem alteração para cada novo módulo C-Level.",
    "Apenas os enums ExecutiveDomain e FrameworkType precisam ser expandidos por módulo — alteração cirúrgica, sem refatoração estrutural.",
    "O frontend é parametrizado por domínio — o mesmo componente ActionPanel e o stepper de estados servem todos os C-Levels.",
    "A OpenAPI spec evolui por módulo (COO_OpenAPI_v1.yaml, CIO_OpenAPI_v1.yaml etc.), mantendo a mesma estrutura do CFO_OpenAPI_v2.yaml.",
    "O cache Redis e o fallback determinístico operam de forma transparente para todos os módulos.",
    "O CEO Orchestrator é uma camada de API adicional que agrega os endpoints dos módulos — não substitui nenhum deles.",
    "O conjunto de especificações C-Level (OpenAPI YAMLs + este Roadmap) constitui o artefato central de governança do produto.",
]
for p_text in premissas:
    add_bullet(doc, p_text)

doc.add_paragraph()
add_paragraph(doc,
    "Este documento é um artefato vivo — revisado a cada milestone do roadmap.",
    size=10, italic=True, color=(130, 130, 130), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUTPUT)
print(f"✅  Documento gerado com sucesso: {OUTPUT}")
