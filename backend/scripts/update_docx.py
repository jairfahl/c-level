"""
Script para atualizar todos os arquivos .docx do projeto C-Level Mentor
com o estado atual da implementação (v3).
"""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = "/Users/jairfahl/Downloads/c-level"


def make_doc(title, subtitle, sections):
    """Create a .docx with title, subtitle, and sections list.
    Each section is (heading, [paragraphs]) where paragraphs can be str or ('bullet', str).
    """
    doc = docx.Document()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for heading, paragraphs in sections:
        doc.add_heading(heading, level=2)
        for item in paragraphs:
            if isinstance(item, tuple) and item[0] == 'bullet':
                doc.add_paragraph(item[1], style='List Bullet')
            elif isinstance(item, tuple) and item[0] == 'heading3':
                doc.add_heading(item[1], level=3)
            else:
                doc.add_paragraph(str(item))

    return doc


def update_phase0():
    sections = [
        ("1. Objetivo", [
            "Definir propósito institucional, escopo e critérios de validação do MVP do CFO Mentor, alinhado ao APQC Domínio 9 (Manage Financial Resources).",
        ]),
        ("2. Escopo", [
            "Camada de governança cognitiva aplicada a decisões financeiras relevantes.",
            ("bullet", "Cobre: planejamento, reporting, treasury, estrutura de capital e risco financeiro."),
            ("bullet", "Sem execução transacional (sem pagamentos, sem lançamentos contábeis)."),
            ("bullet", "Uso voluntário no MVP; enforcement por impacto implementado (impact_score >= 4)."),
            ("bullet", "Mecanismos anti-terceirização cognitiva ativos para prevenir aceitação passiva de IA."),
        ]),
        ("3. Critérios de Sucesso", [
            "Critérios quantitativos:",
            ("bullet", "Mínimo 5 decisões financeiras históricas processadas. [ATINGIDO: 5/5]"),
            ("bullet", "100% de enforcement do protocolo nos casos simulados. [ATINGIDO: 100%]"),
            ("bullet", "Comparação estruturada entre racional original vs racional com protocolo. [ATINGIDO: 5 Comparison Reports]"),
            ("bullet", "Premissas implícitas capturadas >= 10. [ATINGIDO: 16]"),
            ("bullet", "Teoria dos Jogos ativa em >= 2 casos elegíveis. [ATINGIDO: 2/2]"),
            "KPIs qualitativos [v2]:",
            ("bullet", "Score de clareza da recomendação >= 4/5 em todos os casos."),
            ("bullet", "Ganho cognitivo mensurável: +3.2 premissas/caso em média."),
            ("bullet", "Tempo de estruturação: ~5 min via protocolo vs ~2h em comitê."),
            "KPIs de inteligência decisória [v3]:",
            ("bullet", "Acurácia média de forecast (casos closed)."),
            ("bullet", "Taxa de divergência bem-sucedida (divergence_outcome_flag)."),
            ("bullet", "Taxa de materialização de riscos (risk_realization_rate)."),
            ("bullet", "Eficiência de alocação de capital (capital_allocation_efficiency_score)."),
        ]),
        ("4. Restrições Não-Funcionais", [
            ("bullet", "Latência de análise LLM: <= 30s (timeout configurável)."),
            ("bullet", "Fallback determinístico obrigatório: llm_unavailable=true nunca impede transição."),
            ("bullet", "Cache Redis 24h para respostas LLM idênticas."),
            ("bullet", "Audit trail completo: toda ação gera AuditLog."),
            ("bullet", "JWT sem banco de usuários (MVP): qualquer username aceito."),
            ("bullet", "Upload de documentos: máx 10MB, tipos PDF/DOCX/TXT."),
        ]),
        ("5. Anti-Terceirização Cognitiva [v3 NOVO]", [
            "Mecanismos para prevenir 'obesidade de IA' — aceitação passiva de recomendações sem reflexão:",
            ("bullet", "P1 — Reflexão Pré-Recomendação: CFO registra hipótese (min 30 chars) ANTES de ver a recomendação da IA. Backend persiste initial_hypothesis no registro Decision."),
            ("bullet", "P2 — Detecção de Rubber-Stamping: Frontend calcula Jaccard similarity entre decisão e recomendação. Threshold > 0.70 exibe modal de confirmação."),
            ("bullet", "P4 — Reconhecimento de Alertas: Se há alertas heurísticos de casos similares, checkbox obrigatório antes de decidir."),
        ]),
        ("6. Veredicto de Enforcement", [
            "DECISÃO: GO — Mentor CFO aprovado para enforcement no protocolo decisório financeiro.",
            ("bullet", "Enforcement imediato para decisões >= R$ 2M (impact_score >= 4)."),
            ("bullet", "Recomendado para decisões R$ 500k–R$ 2M (impact_score = 3)."),
        ]),
    ]
    doc = make_doc(
        "CFO – Phase 0: Strategic Definition",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_Spec_Phase_0_Strategic_Definition_v2.docx"))


def update_phase1():
    sections = [
        ("1. Modelo de Arquitetura", [
            "Motor híbrido com componentes complementares:",
            ("bullet", "Engine Determinística: enforça protocolo, regras de compliance, transições de estado, seleção de framework (11 tipos), multi-framework selection (até 4)."),
            ("bullet", "Camada LLM: executa raciocínio financeiro estruturado e análise qualitativa, com injeção de documentos da base de conhecimento."),
            ("bullet", "Módulo de Inteligência: alertas heurísticos, benchmark comparativo, KPIs decisórios consolidados."),
            ("bullet", "Anti-Terceirização Cognitiva: hipótese pré-recomendação, detecção de rubber-stamping, reconhecimento de alertas."),
            "Sem o determinístico → vira chatbot.",
            "Sem o LLM → vira checklist burocrático.",
            "Sem anti-terceirização → vira máquina de rubber-stamping.",
        ]),
        ("2. Módulos Core", [
            ("bullet", "FinancialDecisionCase Manager — CRUD e gerenciamento de ciclo de vida"),
            ("bullet", "State Machine Controller — transições validadas e auditadas (8 estados)"),
            ("bullet", "Financial Framework Orchestrator — seleciona e aplica 11 frameworks analíticos (PDCA, cenários, risco, capital, trade-off, game theory, decision matrix, cost-benefit, decision tree, SWOT, Delphi)"),
            ("bullet", "Multi-Framework Selection — permite combinar até 4 frameworks com justificativa automática"),
            ("bullet", "Game Theory Capability — ativada para negociação, precificação competitiva, mercado de capitais"),
            ("bullet", "Audit & Compliance Logger — registra toda ação no sistema (INSERT ONLY)"),
            ("bullet", "Knowledge Base Manager — upload e extração de documentos regulatórios para injeção no contexto LLM"),
            ("bullet", "Intelligence Service — alertas heurísticos, benchmark de casos similares, KPIs decisórios"),
            ("bullet", "Anti-Terceirização Cognitiva — hipótese pré-recomendação, rubber-stamping detection, alert acknowledgment"),
        ]),
        ("3. Fluxo de Dados Entre Módulos [v2]", [
            "Caso criado (DRAFT) → ImpactScorer calcula score → FrameworkSelector escolhe framework → Multi-Framework sugere frameworks adicionais → LLM analisa com Knowledge Base injetada → Hipótese do CFO registrada → Recomendação revelada → Alertas heurísticos exibidos → CFO decide → Review pós-decisão → Heurísticas geradas automaticamente → Dashboard atualizado.",
        ]),
        ("4. Alinhamento APQC Domínio 9", [
            ("bullet", "9.1 Develop and Manage Financial Plans — budget_adjustment, forecast_revision"),
            ("bullet", "9.2 Perform Revenue Accounting — reporting"),
            ("bullet", "9.3 Manage Treasury Operations — treasury, liquidity_management"),
            ("bullet", "9.4 Control and Manage Fixed Assets — capital_allocation, investment_evaluation"),
            ("bullet", "9.5 Process Accounts Payable and Receivable — funding, debt_structuring"),
        ]),
        ("5. Frontend [v3 NOVO]", [
            "Interface web completa em React 18 + Vite + Tailwind CSS + TanStack Query + React Router v6.",
            ("bullet", "Dashboard com lista de casos, criação, filtros por domínio/estado/tipo."),
            ("bullet", "CaseDetail com stepper visual por estado e painéis de ação contextuais."),
            ("bullet", "Multi-Framework Selection UI com checkboxes e rationale."),
            ("bullet", "Gestão de heurísticas, base de conhecimento e administração."),
            ("bullet", "Intelligence Dashboard com 4 KPI cards e tooltips explicativos."),
            ("bullet", "Anti-terceirização cognitiva: hypothesis gate, rubber-stamp modal, alert checkbox."),
        ]),
    ]
    doc = make_doc(
        "CFO – Phase 1: Conceptual Architecture",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_Spec_Phase_1_Conceptual_Architecture_v2.docx"))


def update_phase2():
    sections = [
        ("1. Modelo de Dados – FinancialDecisionCase", [
            "Campos completos do caso decisório:",
            ("bullet", "id: UUID (PK)"),
            ("bullet", "title: string"),
            ("bullet", "description: text"),
            ("bullet", "financial_domain: enum (planning|reporting|treasury|funding|risk)"),
            ("bullet", "impact_score: integer 1–5"),
            ("bullet", "financial_exposure: numeric(18,2)"),
            ("bullet", "time_horizon: enum (short|medium|long) [v2: ENUM corrigido]"),
            ("bullet", "external_agents_present: boolean"),
            ("bullet", "decision_type: enum (8 tipos)"),
            ("bullet", "framework_selected: enum (11 tipos) [v2: adicionado]"),
            ("bullet", "scenario_required: boolean (gerado: impact_score >= 4) [v2: adicionado]"),
            ("bullet", "state: enum DecisionState (8 estados)"),
            ("bullet", "created_at, updated_at: timestamps"),
        ]),
        ("2. Tabelas Relacionadas", [
            ("bullet", "financial_assumptions — premissas (com is_implicit para captura LLM)"),
            ("bullet", "financial_risks — riscos (com materialized para pós-decisão)"),
            ("bullet", "financial_metrics_impacted — métricas impactadas pela decisão"),
            ("bullet", "decisions — recomendação do Mentor + initial_hypothesis [v3] + decisão executiva + divergence_flag"),
            ("bullet", "reviews — resultado real + métricas pós-decisão + divergence_outcome_flag [v2]"),
            ("bullet", "state_transitions — trilha completa de auditoria [v2]"),
            ("bullet", "audit_logs — log imutável INSERT ONLY [v2]"),
            ("bullet", "financial_heuristics — heurísticas aprendidas pós-decisão [v2]"),
            ("bullet", "knowledge_documents — documentos regulatórios para injeção LLM [v3]"),
        ]),
        ("3. Decision Types (8 tipos)", [
            ("bullet", "budget_adjustment → PDCA"),
            ("bullet", "forecast_revision → Análise de Cenários"),
            ("bullet", "capital_allocation → Capital Allocation (Game Theory elegível)"),
            ("bullet", "debt_structuring → Trade-Off (Game Theory elegível)"),
            ("bullet", "liquidity_management → Matriz de Riscos"),
            ("bullet", "risk_hedging → Matriz de Riscos"),
            ("bullet", "cost_reduction → Trade-Off"),
            ("bullet", "investment_evaluation → Análise de Cenários (Game Theory elegível)"),
        ]),
        ("4. Framework Types (11 tipos) [v3 atualizado]", [
            ("bullet", "pdca — Plan-Do-Check-Act"),
            ("bullet", "scenario_analysis — Análise de Cenários (Pessimista/Base/Otimista)"),
            ("bullet", "game_theory — Teoria dos Jogos (players, strategies, payoffs, Nash)"),
            ("bullet", "trade_off — Análise de Trade-Off"),
            ("bullet", "risk_matrix — Matriz de Riscos (probabilidade × impacto)"),
            ("bullet", "capital_allocation — Alocação de Capital"),
            ("bullet", "decision_matrix — Matriz de Decisão (critérios ponderados) [v3]"),
            ("bullet", "cost_benefit_analysis — Análise Custo-Benefício [v3]"),
            ("bullet", "decision_tree — Árvore de Decisão (EMV) [v3]"),
            ("bullet", "swot_analysis — SWOT (Forças, Fraquezas, Oportunidades, Ameaças) [v3]"),
            ("bullet", "delphi_method — Método Delphi (consenso estruturado) [v3]"),
        ]),
        ("5. Máquina de Estados (8 estados)", [
            "DRAFT → CLASSIFIED → STRUCTURED → ANALYZED → RECOMMENDED → DECIDED → UNDER_REVIEW → CLOSED",
            ("bullet", "Cada estado tem exatamente um sucessor permitido."),
            ("bullet", "CLOSED é terminal — nenhuma transição adicional aceita."),
            ("bullet", "Qualquer transição inválida retorna HTTP 409."),
            ("bullet", "Transição DECIDED → UNDER_REVIEW acionada automaticamente após 90 dias (ReviewTrigger)."),
        ]),
        ("6. Regras de Validação do Protocolo", [
            ("bullet", "Mínimo 3 premissas financeiras para estruturar (CLASSIFIED → STRUCTURED)."),
            ("bullet", "Mínimo 3 riscos financeiros para estruturar."),
            ("bullet", "Cenários obrigatórios quando impact_score >= 4."),
            ("bullet", "Game Theory ativa quando external_agents=true E decision_type elegível."),
            ("bullet", "Hipótese pré-recomendação: mínimo 30 caracteres (anti-terceirização cognitiva) [v3]."),
            ("bullet", "Reconhecimento de alertas heurísticos obrigatório antes de decidir (quando presentes) [v3]."),
            ("bullet", "Multi-framework selection: mínimo 1, máximo 4 frameworks por análise [v3]."),
        ]),
        ("7. Knowledge Documents [v3 NOVO]", [
            "Tabela knowledge_documents para upload de documentos regulatórios e políticas internas:",
            ("bullet", "Tipos aceitos: PDF, DOCX, TXT. Limite: 10 MB por arquivo."),
            ("bullet", "Texto extraído automaticamente e armazenado em extracted_text."),
            ("bullet", "Classificado por financial_domain e decision_type."),
            ("bullet", "Injetado no contexto LLM durante análise para enriquecimento contextual."),
            ("bullet", "Soft delete (active=false) — nunca remove fisicamente."),
        ]),
    ]
    doc = make_doc(
        "CFO – Phase 2: Functional Specification",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_Spec_Phase_2_Functional_Specification_v2.docx"))


def update_phase3():
    sections = [
        ("1. Stack Tecnológico", [
            "Stack definido e implementado:",
            ("bullet", "Backend: Python 3.11+ / FastAPI 0.115"),
            ("bullet", "ORM: SQLAlchemy 2.0 (async) + Alembic (migrations)"),
            ("bullet", "Banco: PostgreSQL 15+"),
            ("bullet", "Cache: Redis 7+ (respostas LLM, TTL 24h)"),
            ("bullet", "Auth: JWT via python-jose (HS256)"),
            ("bullet", "LLM: Anthropic Claude (claude-sonnet-4-6)"),
            ("bullet", "Frontend: React 18 + Vite 5 + Tailwind CSS 3 + TanStack Query 5 + React Router 6 [v3]"),
            ("bullet", "Containerização: Docker + Docker Compose"),
        ]),
        ("2. Schema SQL v3", [
            "Schema completo com 10 tabelas:",
            ("bullet", "5 ENUMs: financial_domain (5), decision_state (8), decision_type (8), time_horizon_type (3), framework_type (11)"),
            ("bullet", "financial_decision_cases — entidade central"),
            ("bullet", "financial_assumptions — premissas (com is_implicit)"),
            ("bullet", "financial_risks — riscos (com materialized)"),
            ("bullet", "financial_metrics_impacted — métricas afetadas"),
            ("bullet", "decisions — recomendação + initial_hypothesis [v3] + decisão executiva + divergence_flag"),
            ("bullet", "reviews — resultado pós-decisão + divergence_outcome_flag"),
            ("bullet", "state_transitions — trilha de auditoria de estados"),
            ("bullet", "audit_logs — log imutável (INSERT ONLY via regras SQL)"),
            ("bullet", "financial_heuristics — heurísticas aprendidas"),
            ("bullet", "knowledge_documents — documentos regulatórios [v3]"),
            "Arquivo de referência: CFO/Corrected_v2/CFO_Schema_v2.sql",
        ]),
        ("3. Contrato OpenAPI 3.0 v3", [
            "Contrato completo com 25+ endpoints em 7 grupos:",
            ("bullet", "Auth: POST /auth/token"),
            ("bullet", "Financial Decision Cases: POST, GET list, GET detail, PATCH reclassify"),
            ("bullet", "State Machine: PUT classify|structure|analyze|reanalyze|decide|review|hypothesis, GET suggest-methods|suggest-reclassification|heuristic-alerts"),
            ("bullet", "Audit: GET state-transitions"),
            ("bullet", "Heuristics: POST, GET list, PUT deactivate, GET learning-summary"),
            ("bullet", "Knowledge Base: POST upload, GET list, GET detail, DELETE"),
            ("bullet", "Admin: GET pending-reviews, GET decision-intelligence"),
            "Arquivo de referência: CFO/Corrected_v2/CFO_OpenAPI_v2.yaml",
        ]),
        ("4. Estratégia de Autenticação", [
            ("bullet", "JWT Bearer com HS256."),
            ("bullet", "Endpoint POST /auth/token aceita qualquer username (MVP)."),
            ("bullet", "Token com exp de 8h (configurável via JWT_ACCESS_TOKEN_EXPIRE_MINUTES)."),
            ("bullet", "Sem tabela de usuários — sub do JWT identifica ator para auditoria."),
        ]),
        ("5. Estratégia de Cache", [
            ("bullet", "Redis 7 para cache de respostas LLM."),
            ("bullet", "Chave: SHA256 do user_prompt completo."),
            ("bullet", "TTL: 86400s (24h)."),
            ("bullet", "Falha aberta: se Redis indisponível, chamada LLM prossegue normalmente."),
        ]),
        ("6. Estrutura de Diretórios", [
            "backend/app/api/routers/ — 7 routers (cases, state_machine, audit, heuristics, knowledge_base, admin, auth)",
            "backend/app/core/ — engine determinística (state_machine, impact_scorer, framework_selector, game_theory, audit_logger, auth, config, database, exceptions)",
            "backend/app/llm/ — camada LLM (prompt_builder, client, parser, cache, fallback, service)",
            "backend/app/services/ — serviços (review_service, heuristics_service, intelligence_service, knowledge_base_service, review_trigger)",
            "backend/app/models/ — ORM (enums, base, financial_decision_case, knowledge_document)",
            "backend/app/schemas/ — schemas Pydantic v2 (30+ schemas em __init__.py)",
            "backend/tests/ — 378+ testes (unit/ + integration/)",
            "frontend/src/ — React 18 (pages, components, lib)",
        ]),
        ("7. Docker Compose", [
            "4 serviços:",
            ("bullet", "app — FastAPI backend (porta 8000)"),
            ("bullet", "postgres — PostgreSQL 15 (porta 5432)"),
            ("bullet", "redis — Redis 7 (porta 6379)"),
            ("bullet", "frontend — React (porta 3000, proxy /v1/* → app:8000)"),
        ]),
    ]
    doc = make_doc(
        "CFO – Phase 3: Technical Architecture",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_Spec_Phase_3_Technical_Architecture_v2.docx"))


def update_phase4():
    sections = [
        ("1. Responsabilidades da Engine Determinística", [
            ("bullet", "Scoring de impacto financeiro por faixas de exposure (5 faixas: <100k, 100k-500k, 500k-2M, 2M-10M, >10M)."),
            ("bullet", "Enforcement de análise de cenários (scenario_required quando impact_score >= 4)."),
            ("bullet", "Validação de transições de estado (8 estados, grafo linear)."),
            ("bullet", "Ativação da Teoria dos Jogos por critério formal (external_agents + decision_type elegível)."),
            ("bullet", "Seleção automática de framework por decision_type (11 tipos)."),
            ("bullet", "Multi-Framework Selection: sugestão automática de frameworks complementares com rationale em PT-BR."),
            ("bullet", "Registro de divergência (divergence_flag) com justificativa e critérios de monitoramento."),
        ]),
        ("2. Responsabilidades da Camada LLM", [
            ("bullet", "Raciocínio financeiro estruturado com prompt System + User."),
            ("bullet", "Análise multi-framework: LLM aplica cada framework como lente complementar e sintetiza recomendação integrada."),
            ("bullet", "Injeção de Knowledge Base: documentos relevantes inseridos no prompt como contexto adicional."),
            ("bullet", "Explicitação de premissas implícitas (is_implicit=true)."),
            ("bullet", "Articulação de cenários de fluxo de caixa (pessimista/base/otimista)."),
            ("bullet", "Modelagem qualitativa de equilíbrio (Game Theory: players, strategies, payoffs, equilibrium)."),
            ("bullet", "Sugestão de reclassificação de domínio/tipo de decisão."),
            ("bullet", "Resumo de aprendizado institucional (learning-summary)."),
        ]),
        ("3. Template de Prompt LLM", [
            "SYSTEM PROMPT:",
            '"You are a structured financial decision-making mentor for CFOs. Your role is to analyze financial decisions using rigorous methodology. Never provide opinions without structure. Always identify risks. Never skip assumption explicitation. Output must follow the exact JSON contract defined below."',
            "USER PROMPT sections:",
            ("bullet", "## Context (decision_type, domain, exposure, time_horizon, frameworks, scenario_required, game_theory_active)"),
            ("bullet", "## Knowledge Base (documentos regulatórios relevantes injetados)"),
            ("bullet", "## Stated Assumptions (lista de premissas declaradas)"),
            ("bullet", "## Identified Risks (lista de riscos identificados)"),
            ("bullet", "## Instructions (numbered, referenciando cada framework selecionado)"),
            ("bullet", "## Output JSON Contract (schema exato esperado)"),
        ]),
        ("4. Faixas de Impacto", [
            ("bullet", "Score 1 (Baixo): exposure < R$ 100.000"),
            ("bullet", "Score 2 (Moderado): R$ 100.000 – R$ 500.000"),
            ("bullet", "Score 3 (Relevante): R$ 500.000 – R$ 2.000.000"),
            ("bullet", "Score 4 (Alto): R$ 2.000.000 – R$ 10.000.000 [scenario_required=true]"),
            ("bullet", "Score 5 (Crítico): > R$ 10.000.000 [scenario_required=true]"),
        ]),
        ("5. Framework Selection (11 tipos)", [
            "Mapeamento base decision_type → framework_type:",
            ("bullet", "budget_adjustment → pdca"),
            ("bullet", "forecast_revision → scenario_analysis"),
            ("bullet", "capital_allocation → capital_allocation (Game Theory elegível)"),
            ("bullet", "debt_structuring → trade_off (Game Theory elegível)"),
            ("bullet", "liquidity_management → risk_matrix"),
            ("bullet", "risk_hedging → risk_matrix"),
            ("bullet", "cost_reduction → trade_off"),
            ("bullet", "investment_evaluation → scenario_analysis (Game Theory elegível)"),
            "Sobrescrita: external_agents=true + tipo elegível → game_theory.",
            "Frameworks complementares [v3]: decision_matrix, cost_benefit_analysis, decision_tree, swot_analysis, delphi_method — disponíveis para seleção pelo executivo (máx 4 total).",
        ]),
        ("6. Audit Logger", [
            "AuditAction constants:",
            ("bullet", "CASE_CREATED — caso criado"),
            ("bullet", "STATE_TRANSITION — transição de estado"),
            ("bullet", "LLM_CALLED — chamada LLM bem-sucedida"),
            ("bullet", "LLM_FALLBACK — fallback determinístico acionado"),
            ("bullet", "DIVERGENCE_RECORDED — divergência do CFO registrada"),
            ("bullet", "HEURISTIC_GENERATED — heurística gerada automaticamente"),
            ("bullet", "DOCUMENT_UPLOADED — documento uploaded na knowledge base"),
            ("bullet", "DOCUMENT_DELETED — documento soft-deleted"),
        ]),
        ("7. Knowledge Base [v3 NOVO]", [
            "Módulo de upload e gestão de documentos regulatórios:",
            ("bullet", "Upload via multipart/form-data (PDF, DOCX, TXT, máx 10MB)."),
            ("bullet", "Extração automática de texto (python-docx, PyPDF2, plaintext)."),
            ("bullet", "Classificação por financial_domain e decision_type."),
            ("bullet", "Injeção automática no prompt LLM durante análise."),
            ("bullet", "Soft delete (active=false)."),
        ]),
        ("8. Inteligência Decisória [v3 NOVO]", [
            ("bullet", "Alertas heurísticos: padrões de casos similares exibidos antes da decisão."),
            ("bullet", "Benchmark comparativo: estatísticas de casos similares (acurácia, divergência, eficiência)."),
            ("bullet", "Dashboard KPIs: acurácia de forecast, divergência, materialização de riscos, eficiência de capital."),
            ("bullet", "Performance por domínio financeiro."),
            ("bullet", "Resumo de aprendizado institucional (LLM-generated ou determinístico)."),
        ]),
        ("9. Anti-Terceirização Cognitiva [v3 NOVO]", [
            "3 mecanismos de fricção reflexiva:",
            ("bullet", "P1 — Hipótese Pré-Recomendação: endpoint PUT /hypothesis grava initial_hypothesis na tabela decisions. Frontend oculta recomendação até hipótese registrada."),
            ("bullet", "P2 — Detecção de Rubber-Stamping: frontend calcula Jaccard similarity (normalização: lowercase, strip accents, remove pontuação, filter palavras <= 2 chars). Threshold > 0.70 → modal de confirmação."),
            ("bullet", "P4 — Reconhecimento de Alertas: frontend verifica heuristic-alerts. Se há alertas, checkbox obrigatório antes do botão decidir."),
        ]),
        ("10. Cobertura de Testes", [
            "378+ testes passando, 99% de cobertura:",
            ("bullet", "test_schemas.py — 73 testes (validações Pydantic v2)"),
            ("bullet", "test_engine.py — 108 testes (state machine, scorer, selector, game theory)"),
            ("bullet", "test_llm.py — 76 testes (prompt builder, parser, cache, fallback, service)"),
            ("bullet", "test_learning.py — 47 testes (review service, heuristics, trigger)"),
            ("bullet", "test_api.py — 54 testes (endpoints REST, incluindo hypothesis)"),
            ("bullet", "test_flow.py — 20 testes (fluxo end-to-end DRAFT → CLOSED)"),
        ]),
    ]
    doc = make_doc(
        "CFO – Phase 4: Core Engine Development",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_Spec_Phase_4_Core_Engine_Development_v2.docx"))


def update_phase5():
    sections = [
        ("1. Casos Históricos para Simulação", [
            "5 casos documentados com campos completos, premissas, riscos e resultado real:",
            ("bullet", "Caso 1 — Realocação Orçamentária Q2: COGS vs. Marketing (planning, budget_adjustment, R$ 380k, score 2)"),
            ("bullet", "Caso 2 — Renegociação de CCBs: Sindicato de Bancos R$ 52M (funding, debt_structuring, score 5, Game Theory ativo)"),
            ("bullet", "Caso 3 — Gestão de Liquidez: Gap de Caixa R$ 3,8M (treasury, liquidity_management, score 4)"),
            ("bullet", "Caso 4 — CAPEX Nova Planta Recife R$ 18M (funding, investment_evaluation, score 5, Game Theory ativo)"),
            ("bullet", "Caso 5 — Revisão Forecast Q3: Impacto Cambial (reporting, forecast_revision, R$ 1,2M, score 3)"),
        ]),
        ("2. Resultado da Simulação", [
            ("bullet", "5/5 casos com análise LLM real (Claude Sonnet, 0 fallbacks)."),
            ("bullet", "100% alinhamento CFO/Mentor."),
            ("bullet", "22 premissas declaradas + 16 premissas implícitas capturadas (42%)."),
            ("bullet", "22 riscos estruturados."),
            ("bullet", "4/4 cenários gerados em casos com impact >= 4."),
            ("bullet", "2/2 casos elegíveis com Teoria dos Jogos ativa."),
            ("bullet", "0 divergências registradas na simulação."),
        ]),
        ("3. Heurísticas Aprendidas", [
            "3 heurísticas extraídas dos 5 casos simulados:",
            ("bullet", "bilateral_negotiation_over_syndicate (debt_structuring + >= 3 credores, confiança 0.87) — Caso 02"),
            ("bullet", "hybrid_coverage_for_short_liquidity_gaps (liquidity_management + short + R$1–10M, confiança 0.91) — Caso 03"),
            ("bullet", "pe_fund_alignment_in_coinvestment_structures (investment_evaluation + PE co-investidor, confiança 0.83) — Caso 04"),
        ]),
        ("4. Ganho Cognitivo Mensurável", [
            ("bullet", "+3.2 premissas implícitas por caso em média."),
            ("bullet", "Riscos priorizados por severidade estruturada (vs. lista plana do CFO)."),
            ("bullet", "+3 cenários formalizados por caso crítico."),
            ("bullet", "Tempo de análise: ~5 min (protocolo) vs ~2h (comitê)."),
            ("bullet", "Trilha de auditoria completa em banco (vs. email/PowerPoint)."),
        ]),
        ("5. Comparison Report Template", [
            "Cada caso gera um Comparison Report com 3 seções:",
            ("bullet", "Seção A — Racional Original: justificativa original, premissas e riscos do CFO."),
            ("bullet", "Seção B — Racional Estruturado: framework aplicado, premissas explícitas + implícitas, riscos priorizados, cenários, recomendação."),
            ("bullet", "Seção C — Delta de Qualidade: premissas_adicionais, riscos_adicionais, cenarios_gerados, divergencia, score_clareza_executivo."),
        ]),
        ("6. Veredicto", [
            "DECISÃO: GO — Todos os 5 critérios de qualidade atendidos.",
            ("bullet", "Enforcement imediato para exposições >= R$ 2M."),
            ("bullet", "Recomendado para exposições R$ 500k–R$ 2M."),
        ]),
    ]
    doc = make_doc(
        "CFO – Phase 5: MVP Simulation",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_Spec_Phase_5_MVP_Simulation_v2.docx"))


def update_phase6():
    sections = [
        ("1. Lógica de Revisão Pós-Decisão", [
            ("bullet", "Comparar performance financeira projetada vs realizada."),
            ("bullet", "Capturar riscos financeiros materializados (risk_realization_rate > 50% → materialized=true)."),
            ("bullet", "Calcular divergence_outcome_flag automaticamente (divergence_flag=true AND forecast_accuracy_score < 5)."),
            ("bullet", "Gerar heurísticas automaticamente a partir dos resultados [v3]."),
            ("bullet", "Atualizar dashboard de inteligência decisória [v3]."),
        ]),
        ("2. Trigger de Review", [
            ("bullet", "Casos DECIDED há > 90 dias sem review são listados no endpoint GET /admin/pending-reviews."),
            ("bullet", "Threshold de 90 dias configurável."),
            ("bullet", "No MVP: alerta manual via dashboard. Futuro: notificação automática via email/webhook."),
        ]),
        ("3. Heuristics Registry", [
            ("bullet", "Tabela financial_heuristics com UNIQUE(decision_type, financial_domain, heuristic_key)."),
            ("bullet", "Valor em JSONB (flexível para diferentes padrões aprendidos)."),
            ("bullet", "Source_case_id opcional — vincula à evidência original."),
            ("bullet", "INSERT/UPDATE only — nunca DELETE (apenas deactivate, active=false)."),
            ("bullet", "Criação manual via POST /heuristics (MVP) + geração automática no review [v3]."),
        ]),
        ("4. Alertas Heurísticos [v3 NOVO]", [
            "Quando um caso atinge RECOMMENDED, o Mentor consulta heurísticas de casos similares:",
            ("bullet", "Endpoint GET /financial-decision-cases/{id}/heuristic-alerts."),
            ("bullet", "Filtra heurísticas ativas por decision_type + financial_domain."),
            ("bullet", "Retorna alertas com tipo, severidade, mensagem e confiança."),
            ("bullet", "Frontend exige reconhecimento (checkbox) antes de permitir decisão."),
        ]),
        ("5. Benchmark Comparativo [v3 NOVO]", [
            "Junto com os alertas, retorna estatísticas de casos similares encerrados:",
            ("bullet", "total_similar_cases — quantos casos com mesmo tipo+domínio foram closed."),
            ("bullet", "followed_recommendation_count/pct — % que seguiu a recomendação."),
            ("bullet", "diverged_count + diverged_success_rate — taxa de sucesso das divergências."),
            ("bullet", "avg_forecast_accuracy, avg_risk_realization, avg_capital_efficiency — médias."),
            ("bullet", "most_effective_framework — framework com melhor resultado."),
        ]),
        ("6. Intelligence Dashboard [v3 NOVO]", [
            "Endpoint GET /admin/decision-intelligence retorna KPIs consolidados:",
            ("bullet", "total_cases_closed + total_cases_active."),
            ("bullet", "avg_forecast_accuracy (1–10)."),
            ("bullet", "divergence_total + divergence_success_count + divergence_success_rate."),
            ("bullet", "avg_risk_realization (0–100%)."),
            ("bullet", "avg_capital_efficiency (0–100%)."),
            ("bullet", "domain_performance — breakdown por domínio financeiro."),
            ("bullet", "top_frameworks — frameworks mais utilizados com contagem."),
        ]),
        ("7. Learning Summary [v3 NOVO]", [
            "Endpoint GET /heuristics/learning-summary retorna resumo institucional:",
            ("bullet", "summary — texto resumindo o aprendizado (LLM-generated quando disponível)."),
            ("bullet", "heuristics_count — total de heurísticas ativas."),
            ("bullet", "top_frameworks — frameworks mais utilizados."),
            ("bullet", "last_updated — data da última heurística criada."),
            ("bullet", "llm_generated — se o resumo foi gerado via LLM ou determinístico."),
        ]),
    ]
    doc = make_doc(
        "CFO – Phase 6: Learning Module",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_Spec_Phase_6_Learning_Module_v2.docx"))


def update_ddl_spec():
    sections = [
        ("1. ENUM Types", [
            "5 ENUMs definidos:",
            ("bullet", "financial_domain: planning, reporting, treasury, funding, risk"),
            ("bullet", "decision_state: DRAFT, CLASSIFIED, STRUCTURED, ANALYZED, RECOMMENDED, DECIDED, UNDER_REVIEW, CLOSED"),
            ("bullet", "decision_type: budget_adjustment, forecast_revision, capital_allocation, debt_structuring, liquidity_management, risk_hedging, cost_reduction, investment_evaluation"),
            ("bullet", "time_horizon_type: short, medium, long [v2: ENUM corrigido de VARCHAR(10)]"),
            ("bullet", "framework_type: pdca, scenario_analysis, game_theory, trade_off, risk_matrix, capital_allocation, decision_matrix, cost_benefit_analysis, decision_tree, swot_analysis, delphi_method [v3: 5 frameworks adicionais]"),
        ]),
        ("2. Tabela Principal: financial_decision_cases", [
            ("bullet", "id: UUID PK (gen_random_uuid)"),
            ("bullet", "title, description: TEXT NOT NULL"),
            ("bullet", "financial_domain: financial_domain ENUM NOT NULL"),
            ("bullet", "impact_score: INT CHECK (1-5)"),
            ("bullet", "financial_exposure: NUMERIC(18,2) NOT NULL CHECK (> 0)"),
            ("bullet", "time_horizon: time_horizon_type ENUM"),
            ("bullet", "external_agents_present: BOOLEAN DEFAULT FALSE"),
            ("bullet", "decision_type: decision_type ENUM NOT NULL"),
            ("bullet", "framework_selected: framework_type ENUM [v2]"),
            ("bullet", "scenario_required: BOOLEAN GENERATED ALWAYS AS (impact_score >= 4) STORED [v2]"),
            ("bullet", "state: decision_state ENUM DEFAULT 'DRAFT'"),
            ("bullet", "created_at, updated_at: TIMESTAMP DEFAULT NOW()"),
        ]),
        ("3. Tabelas de Relacionamento", [
            ("bullet", "financial_assumptions — id, decision_case_id (FK CASCADE), text, is_implicit, created_at"),
            ("bullet", "financial_risks — id, decision_case_id (FK CASCADE), text, materialized, created_at"),
            ("bullet", "financial_metrics_impacted — id, decision_case_id (FK CASCADE), metric_name, created_at"),
        ]),
        ("4. Tabela: decisions", [
            ("bullet", "id: UUID PK"),
            ("bullet", "decision_case_id: UUID FK CASCADE"),
            ("bullet", "recommendation: TEXT NOT NULL"),
            ("bullet", "initial_hypothesis: TEXT [v3 — hipótese do CFO antes de ver a recomendação]"),
            ("bullet", "executive_decision: TEXT"),
            ("bullet", "divergence_flag: BOOLEAN DEFAULT FALSE"),
            ("bullet", "created_at: TIMESTAMP"),
        ]),
        ("5. Tabela: reviews", [
            ("bullet", "outcome_summary: TEXT"),
            ("bullet", "forecast_accuracy_score: INT CHECK (1-10)"),
            ("bullet", "risk_realization_rate: NUMERIC(5,2) CHECK (0-100)"),
            ("bullet", "capital_allocation_efficiency_score: NUMERIC(5,2) CHECK (0-100)"),
            ("bullet", "divergence_outcome_flag: BOOLEAN DEFAULT FALSE [v2]"),
        ]),
        ("6. Tabelas de Auditoria [v2]", [
            ("bullet", "state_transitions — from_state, to_state, transitioned_at, triggered_by"),
            ("bullet", "audit_logs — action TEXT, payload JSONB, created_at (INSERT ONLY via regras SQL)"),
        ]),
        ("7. Tabela: financial_heuristics [v2]", [
            ("bullet", "decision_type + financial_domain + heuristic_key: UNIQUE constraint"),
            ("bullet", "heuristic_value: JSONB"),
            ("bullet", "source_case_id: UUID FK SET NULL"),
            ("bullet", "active: BOOLEAN DEFAULT TRUE"),
        ]),
        ("8. Tabela: knowledge_documents [v3 NOVO]", [
            ("bullet", "id: UUID PK"),
            ("bullet", "title, description, original_filename: TEXT"),
            ("bullet", "file_type: VARCHAR(10) CHECK IN ('pdf', 'docx', 'txt')"),
            ("bullet", "file_size_bytes: INT CHECK (> 0 AND <= 10485760)"),
            ("bullet", "extracted_text: TEXT NOT NULL"),
            ("bullet", "text_length: INT NOT NULL"),
            ("bullet", "financial_domain: financial_domain ENUM NOT NULL"),
            ("bullet", "decision_type: decision_type ENUM"),
            ("bullet", "active: BOOLEAN DEFAULT TRUE"),
            ("bullet", "uploaded_by: TEXT"),
            ("bullet", "INDEX: idx_kd_domain_type_active ON (financial_domain, decision_type, active)"),
        ]),
        ("9. Triggers e Índices", [
            ("bullet", "update_updated_at() trigger em financial_decision_cases e financial_heuristics."),
            ("bullet", "16 índices em tabelas principais para performance de queries."),
            ("bullet", "audit_logs_no_update e audit_logs_no_delete rules para imutabilidade."),
        ]),
        ("10. Referência", [
            "Arquivo SQL completo: CFO/Corrected_v2/CFO_Schema_v2.sql",
            "Migrations Alembic: backend/alembic/versions/ (001, 002, 003)",
        ]),
    ]
    doc = make_doc(
        "CFO – Database DDL Specification",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026  |  PostgreSQL 15+",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_DDL_SQL_Spec_v2.docx"))


def update_openapi_spec():
    sections = [
        ("1. Informações Gerais", [
            ("bullet", "OpenAPI 3.0.3"),
            ("bullet", "Versão da API: 3.0.0"),
            ("bullet", "Base URL: https://api.mentor-cfo.internal/v1"),
            ("bullet", "Autenticação: Bearer JWT (HS256) em todos os endpoints (exceto /auth/token)"),
        ]),
        ("2. Endpoints — Auth", [
            ("bullet", "POST /auth/token — Obter token JWT (qualquer username no MVP)"),
        ]),
        ("3. Endpoints — Financial Decision Cases", [
            ("bullet", "POST /financial-decision-cases — Criar caso (→ DRAFT)"),
            ("bullet", "GET /financial-decision-cases — Listar (paginado, filtros: domain, state, decision_type)"),
            ("bullet", "GET /financial-decision-cases/{id} — Detalhe completo (com assumptions, risks, recommendation, initial_hypothesis, executive_decision)"),
            ("bullet", "PATCH /financial-decision-cases/{id}/reclassify — Reclassificar domínio e tipo"),
        ]),
        ("4. Endpoints — State Machine", [
            ("bullet", "PUT /{id}/classify — DRAFT → CLASSIFIED (impact_score obrigatório)"),
            ("bullet", "PUT /{id}/structure — CLASSIFIED → STRUCTURED (min 3 assumptions, min 3 risks)"),
            ("bullet", "GET /{id}/suggest-methods — Sugestões de framework (com rationale em PT-BR)"),
            ("bullet", "GET /{id}/suggest-reclassification — Sugestão de reclassificação via LLM"),
            ("bullet", "PUT /{id}/analyze — STRUCTURED → RECOMMENDED (LLM + engine, body opcional com frameworks_selected)"),
            ("bullet", "PUT /{id}/reanalyze — RECOMMENDED → RECOMMENDED (reanálise após reclassificação)"),
            ("bullet", "PUT /{id}/hypothesis — Registrar hipótese pré-recomendação (min 30 chars) [v3]"),
            ("bullet", "GET /{id}/heuristic-alerts — Alertas heurísticos + benchmark [v3]"),
            ("bullet", "PUT /{id}/decide — RECOMMENDED → DECIDED (executive_decision, divergence, monitoring_criteria)"),
            ("bullet", "PUT /{id}/review — DECIDED → CLOSED (outcome_summary, métricas, heuristics_generated)"),
        ]),
        ("5. Endpoints — Audit", [
            ("bullet", "GET /{id}/state-transitions — Trilha completa de auditoria"),
        ]),
        ("6. Endpoints — Heuristics", [
            ("bullet", "POST /heuristics — Criar heurística manual"),
            ("bullet", "GET /heuristics — Listar (filtros: decision_type, domain)"),
            ("bullet", "PUT /heuristics/{id}/deactivate — Soft delete"),
            ("bullet", "GET /heuristics/learning-summary — Resumo de aprendizado institucional [v3]"),
        ]),
        ("7. Endpoints — Knowledge Base [v3 NOVO]", [
            ("bullet", "POST /knowledge-base/upload — Upload multipart/form-data (file, title, financial_domain)"),
            ("bullet", "GET /knowledge-base — Listar documentos (filtros: domain, decision_type)"),
            ("bullet", "GET /knowledge-base/{id} — Detalhe com texto extraído"),
            ("bullet", "DELETE /knowledge-base/{id} — Soft delete"),
        ]),
        ("8. Endpoints — Admin [v3 NOVO]", [
            ("bullet", "GET /admin/pending-reviews — Casos DECIDED > 90 dias sem review"),
            ("bullet", "GET /admin/decision-intelligence — Dashboard KPIs decisórios"),
        ]),
        ("9. Schemas Principais", [
            ("bullet", "FinancialDecisionCase, FinancialDecisionCaseFull (com initial_hypothesis)"),
            ("bullet", "Assumption, Risk, StateTransition"),
            ("bullet", "AnalysisResult (com game_theory_model, frameworks_selected)"),
            ("bullet", "MethodSuggestionResponse, MethodSelectionRequest"),
            ("bullet", "HypothesisRequest (min_length 30) [v3]"),
            ("bullet", "HeuristicAlertItem, DecisionBenchmark, HeuristicAlertsResponse [v3]"),
            ("bullet", "HeuristicCreate, HeuristicResponse, LearningSummaryResponse"),
            ("bullet", "KnowledgeDocument schemas (Upload, Summary, Full, List) [v3]"),
            ("bullet", "PendingReviewsResponse, DecisionIntelligenceResponse [v3]"),
            ("bullet", "ErrorResponse (error + message + optional details)"),
        ]),
        ("10. Referência", [
            "Contrato YAML completo: CFO/Corrected_v2/CFO_OpenAPI_v2.yaml",
            "Total: 25+ endpoints, 30+ schemas, 7 tags.",
        ]),
    ]
    doc = make_doc(
        "CFO – OpenAPI Contract Specification",
        "Versão 3.0  |  Mentor C-Level – CFO  |  Março 2026  |  OpenAPI 3.0",
        sections,
    )
    doc.save(os.path.join(BASE, "CFO/Corrected_v2/CFO_OpenAPI_Spec_v2.docx"))


def update_roadmap():
    sections = [
        ("Posicionamento do Produto", [
            "O C-Level Mentor não é um chatbot nem um assistente conversacional. É uma plataforma de agentes especializados — sistema que combina engine determinística + LLM para governança cognitiva de decisões executivas.",
            "O MVP CFO está 100% implementado e validado com 5 casos históricos reais. O Mentor CFO impõe um protocolo decisório mandatório com 8 estados, 11 frameworks analíticos, base de conhecimento injetada no LLM, aprendizado institucional via heurísticas, inteligência decisória com KPIs consolidados e mecanismos anti-terceirização cognitiva.",
        ]),
        ("Os componentes de cada agente C-Level", [
            ("heading3", "Guardrails — Defesa em Camadas"),
            ("bullet", "State machine linear — impede transições inválidas (HTTP 409); o agente nunca pula etapas."),
            ("bullet", "Fallback determinístico — qualquer falha do LLM aciona análise estruturada; o agente nunca para."),
            ("bullet", "Audit trail completo — cada ação registrada com timestamp, ator e payload."),
            ("bullet", "Intervenção humana obrigatória — o agente recomenda; o executivo decide (RECOMMENDED → DECIDED)."),
            ("bullet", "Anti-terceirização cognitiva — hipótese pré-recomendação, rubber-stamp detection, alert acknowledgment."),
            ("bullet", "Revisão pós-decisão — ciclo de aprendizado estruturado (UNDER_REVIEW → CLOSED)."),
            ("heading3", "Protocolo Decisório"),
            ("bullet", "8 estados obrigatórios: DRAFT → CLASSIFIED → STRUCTURED → ANALYZED → RECOMMENDED → DECIDED → UNDER_REVIEW → CLOSED."),
            ("bullet", "Mínimo 3 premissas + 3 riscos para estruturar."),
            ("bullet", "Cenários obrigatórios para impact_score >= 4."),
            ("bullet", "Multi-framework selection: até 4 frameworks combinados por análise."),
            ("heading3", "Inteligência Institucional"),
            ("bullet", "Heurísticas aprendidas automaticamente de casos encerrados."),
            ("bullet", "Alertas heurísticos de casos similares exibidos antes da decisão."),
            ("bullet", "Benchmark comparativo com estatísticas de performance."),
            ("bullet", "Dashboard de KPIs: acurácia, divergência, materialização, eficiência."),
        ]),
        ("CFO Mentor — Status: 100% Implementado", [
            "Stack: FastAPI + PostgreSQL + Redis + React 18 + Docker Compose",
            ("bullet", "25+ endpoints REST (auth, cases, state machine, heuristics, knowledge base, admin)"),
            ("bullet", "11 frameworks analíticos (PDCA, Cenários, Game Theory, Trade-Off, Risk Matrix, Capital Allocation, Decision Matrix, Cost-Benefit, Decision Tree, SWOT, Delphi)"),
            ("bullet", "Knowledge Base: upload de documentos regulatórios injetados no contexto LLM"),
            ("bullet", "Anti-Terceirização Cognitiva: hipótese, rubber-stamping, alert acknowledgment"),
            ("bullet", "378+ testes, 99% cobertura"),
            ("bullet", "5 casos MVP validados — Veredicto: GO para enforcement"),
        ]),
        ("CEO Orchestrator — Planejado para Q2–Q3 2027", [
            "O CEO Orchestrator é o agente-mestre que consolida inputs dos outros C-Level agents (CFO, COO, CIO, CMO, CRO) para decisões estratégicas cross-funcionais.",
            ("bullet", "Consolida recomendações de múltiplos agentes especializados."),
            ("bullet", "Detecta conflitos entre domínios (ex: CFO recomenda corte, CMO recomenda investimento)."),
            ("bullet", "Aplica trade-off estruturado com framework cross-domínio."),
            ("bullet", "Mantém visão unificada do portfolio decisório executivo."),
        ]),
        ("Roadmap de Agentes C-Level", [
            ("bullet", "CFO — Governança de decisões financeiras [IMPLEMENTADO ✅]"),
            ("bullet", "COO — Eficiência operacional: otimização de processos, capacidade produtiva, supply chain, KPIs operacionais. [PLANEJADO Q3 2026]"),
            ("bullet", "CIO — Governança de tecnologia: CAPEX tecnológico, cyber-risco, transformação digital, ROI de projetos. [PLANEJADO Q4 2026]"),
            ("bullet", "CMO — Governança de mercado: ROI de campanhas, pricing, market share, brand equity. [PLANEJADO Q1 2027]"),
            ("bullet", "CRO — Governança de receita: pipeline comercial, churn, upsell, previsibilidade de receita. [PLANEJADO Q1 2027]"),
            ("bullet", "CEO Orchestrator — Agente-mestre cross-funcional. [PLANEJADO Q2–Q3 2027]"),
        ]),
        ("Princípios de Escalabilidade", [
            ("bullet", "Cada agente C-Level segue a mesma arquitetura: engine determinística + LLM + state machine + audit trail."),
            ("bullet", "Framework comum: o CFO Mentor define o padrão que será replicado para os demais papéis."),
            ("bullet", "Dados compartilhados: heurísticas e knowledge base são específicas por domínio, mas o CEO Orchestrator tem visão consolidada."),
            ("bullet", "Anti-terceirização cognitiva é um princípio transversal: todos os agentes devem implementar mecanismos de fricção reflexiva."),
        ]),
    ]
    doc = make_doc(
        "C-LEVEL MENTOR — Roadmap de Produto",
        "Versão 2.0  |  Março 2026  |  Visão de Futuro e Expansão para o C-Suite",
        sections,
    )
    doc.save(os.path.join(BASE, "C-Level_Mentor_Roadmap.docx"))


if __name__ == "__main__":
    update_phase0()
    print("✓ Phase 0 updated")
    update_phase1()
    print("✓ Phase 1 updated")
    update_phase2()
    print("✓ Phase 2 updated")
    update_phase3()
    print("✓ Phase 3 updated")
    update_phase4()
    print("✓ Phase 4 updated")
    update_phase5()
    print("✓ Phase 5 updated")
    update_phase6()
    print("✓ Phase 6 updated")
    update_ddl_spec()
    print("✓ DDL Spec updated")
    update_openapi_spec()
    print("✓ OpenAPI Spec updated")
    update_roadmap()
    print("✓ Roadmap updated")
    print("\nAll .docx files updated successfully!")
